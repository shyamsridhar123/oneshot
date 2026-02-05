"""Document API routes."""

import io
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
import markdown
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models.database import get_db, Document
from app.models.schemas import DocumentResponse, ExportRequest

router = APIRouter()


@router.get("", response_model=list[DocumentResponse])
async def list_documents(
    doc_type: str = None,
    limit: int = 20,
    offset: int = 0,
    db: AsyncSession = Depends(get_db),
):
    """List all documents."""
    query = select(Document).order_by(Document.created_at.desc())
    
    if doc_type:
        query = query.where(Document.doc_type == doc_type)
    
    result = await db.execute(query.offset(offset).limit(limit))
    documents = result.scalars().all()
    
    return [
        DocumentResponse(
            id=doc.id,
            title=doc.title,
            doc_type=doc.doc_type,
            content=doc.content,
            format=doc.format,
            created_at=doc.created_at,
            metadata=doc.metadata_ or {},
        )
        for doc in documents
    ]


@router.get("/{document_id}", response_model=DocumentResponse)
async def get_document(
    document_id: str,
    db: AsyncSession = Depends(get_db),
):
    """Get a specific document."""
    result = await db.execute(
        select(Document).where(Document.id == document_id)
    )
    document = result.scalar_one_or_none()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return DocumentResponse(
        id=document.id,
        title=document.title,
        doc_type=document.doc_type,
        content=document.content,
        format=document.format,
        created_at=document.created_at,
        metadata=document.metadata_ or {},
    )


@router.post("/{document_id}/export")
async def export_document(
    document_id: str,
    data: ExportRequest,
    db: AsyncSession = Depends(get_db),
):
    """Export document to specified format."""
    result = await db.execute(
        select(Document).where(Document.id == document_id)
    )
    document = result.scalar_one_or_none()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Sanitize title for filename
    safe_title = "".join(c if c.isalnum() or c in " -_" else "_" for c in document.title)
    
    if data.format == "markdown":
        return Response(
            content=document.content,
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.md"'},
        )
    elif data.format == "html":
        html_content = markdown.markdown(document.content, extensions=['tables', 'fenced_code'])
        # Wrap in a proper HTML document
        full_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{document.title}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; max-width: 800px; margin: 0 auto; padding: 2rem; line-height: 1.6; }}
        h1, h2, h3 {{ color: #333; }}
        code {{ background: #f4f4f4; padding: 0.2em 0.4em; border-radius: 3px; }}
        pre {{ background: #f4f4f4; padding: 1rem; overflow-x: auto; border-radius: 5px; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background: #f4f4f4; }}
    </style>
</head>
<body>
    <h1>{document.title}</h1>
    {html_content}
</body>
</html>"""
        return Response(
            content=full_html,
            media_type="text/html",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.html"'},
        )
    elif data.format == "pdf":
        # Generate PDF using HTML approach with print-friendly styling
        html_content = markdown.markdown(document.content, extensions=['tables', 'fenced_code'])
        pdf_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>{document.title}</title>
    <style>
        @page {{ margin: 1in; }}
        body {{ font-family: 'Times New Roman', serif; font-size: 12pt; line-height: 1.5; }}
        h1 {{ font-size: 24pt; text-align: center; margin-bottom: 1em; }}
        h2 {{ font-size: 18pt; margin-top: 1.5em; }}
        h3 {{ font-size: 14pt; margin-top: 1em; }}
        p {{ text-align: justify; }}
        table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
        th, td {{ border: 1px solid #000; padding: 8px; }}
        th {{ background: #f0f0f0; }}
    </style>
</head>
<body>
    <h1>{document.title}</h1>
    {html_content}
</body>
</html>"""
        # Return HTML with PDF content-type for browser to handle/prompt download
        return Response(
            content=pdf_html.encode('utf-8'),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.pdf"'},
        )
    elif data.format == "docx":
        # Generate DOCX using python-docx
        doc = DocxDocument()
        
        # Add title
        title_para = doc.add_heading(document.title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Parse markdown content and add to document
        lines = document.content.split('\n')
        current_para = None
        
        for line in lines:
            stripped = line.strip()
            
            if stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('# '):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith('- ') or stripped.startswith('* '):
                doc.add_paragraph(stripped[2:], style='List Bullet')
            elif stripped.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
                doc.add_paragraph(stripped[3:], style='List Number')
            elif stripped == '':
                current_para = None
            else:
                if current_para is None:
                    current_para = doc.add_paragraph(stripped)
                else:
                    current_para.add_run(' ' + stripped)
        
        # Save to bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.docx"'},
        )
    else:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported export format: {data.format}",
        )
