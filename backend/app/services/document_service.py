"""Document generation service."""

import uuid
from datetime import datetime
from typing import Optional

from sqlalchemy.ext.asyncio import AsyncSession

from app.models.database import Document


class DocumentService:
    """Service for generating and managing documents."""

    async def create_document(
        self,
        db: AsyncSession,
        title: str,
        doc_type: str,
        content: str,
        format: str = "markdown",
        conversation_id: Optional[str] = None,
        metadata: dict = None,
    ) -> Document:
        """Create a new document."""
        doc = Document(
            id=str(uuid.uuid4()),
            title=title,
            doc_type=doc_type,
            content=content,
            format=format,
            conversation_id=conversation_id,
            metadata_=metadata or {},
        )
        db.add(doc)
        await db.flush()
        return doc

    def export_to_markdown(self, document: Document) -> str:
        """Export document as markdown."""
        return document.content

    def export_to_html(self, document: Document) -> str:
        """Export document as HTML."""
        import markdown
        return markdown.markdown(
            document.content,
            extensions=["tables", "fenced_code"],
        )

    async def export_to_docx(self, document: Document) -> bytes:
        """Export document as DOCX."""
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt
        from io import BytesIO

        doc = DocxDocument()

        # Add title
        doc.add_heading(document.title, 0)

        # Parse markdown and add content
        # Simple conversion - for full markdown support, use a library
        lines = document.content.split("\n")
        for line in lines:
            if line.startswith("# "):
                doc.add_heading(line[2:], 1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], 2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], 3)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.strip():
                doc.add_paragraph(line)

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()


# Singleton
_document_service: DocumentService | None = None


def get_document_service() -> DocumentService:
    """Get or create document service singleton."""
    global _document_service
    if _document_service is None:
        _document_service = DocumentService()
    return _document_service
